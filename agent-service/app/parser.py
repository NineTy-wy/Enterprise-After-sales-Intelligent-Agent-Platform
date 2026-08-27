import csv
import io
from pathlib import Path

from .config import get_settings
from .storage import read_object


def parse_document(file_path: str, file_type: str, supplied_content: str | None = None) -> str:
    """按文件类型解析文本；依赖未安装时返回明确降级提示。"""
    if supplied_content:
        return supplied_content

    path = Path(file_path)
    if not path.exists() and not path.is_absolute():
        path = Path(get_settings().storage_root) / path
    suffix = file_type.lower().lstrip(".") or path.suffix.lower().lstrip(".")
    try:
        if suffix == "pdf":
            from pypdf import PdfReader
            return "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
        if suffix == "docx":
            from docx import Document
            return "\n".join(paragraph.text for paragraph in Document(path).paragraphs)
        if suffix == "doc":
            raise ValueError("暂不支持旧版 .doc 文件，请转换为 .docx 后重新上传")
        if suffix == "xlsx":
            from openpyxl import load_workbook
            workbook = load_workbook(path, read_only=True, data_only=True)
            rows = []
            for sheet in workbook.worksheets:
                rows.append(f"[Sheet: {sheet.title}]")
                rows.extend(" | ".join("" if value is None else str(value)
                            for value in row)
                        for row in sheet.iter_rows(values_only=True))
            return "\n".join(rows)
        if suffix == "xls":
            return _parse_xls_path(path)
        if suffix == "csv":
            with path.open("r", encoding="utf-8-sig", errors="ignore",
                           newline="") as file:
                return "\n".join(
                    " | ".join(cell.strip() for cell in row)
                    for row in csv.reader(file)
                )
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        raise ValueError(f"document parsing failed: {exc}") from exc


def parse_document_bytes(data: bytes, file_type: str) -> str:
    """解析来自 MinIO 的二进制对象。"""
    import io

    suffix = file_type.lower().lstrip(".")
    try:
        if suffix == "pdf":
            from pypdf import PdfReader
            return "\n".join(page.extract_text() or ""
                            for page in PdfReader(io.BytesIO(data)).pages)
        if suffix == "docx":
            from docx import Document
            return "\n".join(paragraph.text
                            for paragraph in Document(io.BytesIO(data)).paragraphs)
        if suffix == "doc":
            raise ValueError("暂不支持旧版 .doc 文件，请转换为 .docx 后重新上传")
        if suffix == "xlsx":
            from openpyxl import load_workbook
            workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            rows = []
            for sheet in workbook.worksheets:
                rows.append(f"[Sheet: {sheet.title}]")
                rows.extend(" | ".join("" if value is None else str(value)
                            for value in row)
                        for row in sheet.iter_rows(values_only=True))
            return "\n".join(rows)
        if suffix == "xls":
            return _parse_xls_bytes(data)
        if suffix == "csv":
            return "\n".join(
                " | ".join(cell.strip() for cell in row)
                for row in csv.reader(
                    io.StringIO(data.decode("utf-8-sig", errors="ignore"))
                )
            )
        return data.decode("utf-8", errors="ignore")
    except Exception as exc:
        raise ValueError(f"document parsing failed: {exc}") from exc


def _parse_xls_path(path: Path) -> str:
    """解析 Excel 97-2003 二进制工作簿。"""
    import xlrd

    workbook = xlrd.open_workbook(filename=str(path), on_demand=True)
    return _workbook_to_text(workbook)


def _parse_xls_bytes(data: bytes) -> str:
    """解析来自 MinIO 的 Excel 97-2003 二进制工作簿。"""
    import xlrd

    workbook = xlrd.open_workbook(file_contents=data, on_demand=True)
    return _workbook_to_text(workbook)


def _workbook_to_text(workbook) -> str:
    rows: list[str] = []
    for sheet in workbook.sheets():
        rows.append(f"[Sheet: {sheet.name}]")
        for row_index in range(sheet.nrows):
            rows.append(" | ".join(
                str(sheet.cell_value(row_index, column_index))
                for column_index in range(sheet.ncols)
            ))
    return "\n".join(rows)
